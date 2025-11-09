# babix_mbft_app.py
# ⚖️ Babix – QQROC + Verificação MBFT + Defesa DOCX
# deps: streamlit, PyPDF2, python-docx

import os, re, io
from difflib import SequenceMatcher
import streamlit as st
import PyPDF2
from docx import Document
from docx.shared import Pt

# ============ PAGE STYLE ============
st.set_page_config(page_title="Babix – QQROC + MBFT", page_icon="⚖️", layout="centered")
st.markdown("""
<style>
:root{--bg:#F6F8FB;--card:#fff;--accent:#0A66C2;--muted:#5b6b7a;--ok:#1b8f3e;--warn:#b88700;--err:#c23b3b;}
.block-container{padding:2rem 2.2rem 3rem}
body{background:var(--bg)}
h1,h2,h3{color:#0b2b4a}
.card{background:var(--card);border-radius:14px;padding:18px;box-shadow:0 10px 26px rgba(16,24,40,.06),0 2px 4px rgba(16,24,40,.03);margin:12px 0}
.badge{display:inline-flex;align-items:center;border-radius:999px;padding:.15rem .6rem;font:700 .8rem system-ui}
.badge.ok{background:#e8f6ed;color:var(--ok)} .badge.warn{background:#fff7e5;color:var(--warn)} .badge.err{background:#fdecec;color:var(--err)}
.obs{background:#f3f7ff;border-left:4px solid var(--accent);padding:.75rem;border-radius:6px}
.code-pill{font-family:ui-monospace,Menlo,monospace;background:#eef2f8;color:#0b2b4a;border-radius:8px;padding:.2rem .45rem}
.kv{display:flex;gap:.5rem;flex-wrap:wrap}
.kv div{background:#f6f8fd;padding:.35rem .6rem;border-radius:8px}
hr{border:none;height:1px;background:#e8eef6;margin:14px 0}
</style>
""", unsafe_allow_html=True)

st.title("⚖️ Babix – Análise de Multas (QQROC + MBFT)")
st.caption("Upload único • Extração robusta • Comparação do campo Observações com a ficha MBFT • Geração de defesa em DOCX")

# ============ HELPERS ============

NA_TITLES = [
    "EMBARCADOR/TRANSPORTADOR",
    "IDENTIFICAÇÃO DO PROPRIETÁRIO",
    "IDENTIFICAÇÃO DO PROPRIETARIO",
    "IDENTIFICAÇÃO DO AGENTE",
    "IDENTIFICAÇÃO DO LOCAL",
    "MENSAGEM SENATRAN",
    "REGISTRO FOTOGRÁFICO",
    "IDENTIFICAÇÃO DO CONDUTOR",
    "IDENTIFICAÇÃO DO VEÍCULO",
    "IDENTIFICAÇÃO DA AUTUAÇÃO",
    "NOTIFICAÇÃO DE AUTUAÇÃO",
    "CÓDIGO DO ÓRGÃO",
    "CÓDIGO DO ÓRGÃO AUTUADOR",
    "CÓDIGO DO MUNICÍPIO",
    "IDENTIFICAÇÃO DA INFRAÇÃO",
]

def read_pdf_text(file_obj) -> str:
    try:
        reader = PyPDF2.PdfReader(file_obj)
        txt = ""
        for p in reader.pages:
            txt += p.extract_text() or ""
        return txt
    except Exception as e:
        st.error(f"Erro ao ler PDF: {e}")
        return ""

def norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

def normalize_codigo(c: str) -> str:
    """Normaliza 52741/5274 1/527-41 -> 527-41"""
    c = (c or "").replace("–", "-").replace(" ", "")
    if "-" in c:
        return c
    # 5 dígitos: 3 + 2 (ex.: 52741)
    if len(c) == 5:
        return f"{c[:3]}-{c[3:]}"
    # 6 dígitos: 4 + 2 (ex.: 705621 -> 7056-21)
    if len(c) == 6:
        return f"{c[:4]}-{c[4:]}"
    # 4+1 com espaço já removido (ex.: 52741 tratado acima)
    return c

def extract_codigo_infracao(texto: str) -> str | None:
    """
    Extrai especificamente o CÓDIGO DA INFRAÇÃO (contextual).
    Aceita: 527-41, 52741, 5274 1 (normaliza p/ 527-41).
    """
    txt = texto.replace("–", "-")
    # forma principal (rótulo explícito)
    m = re.search(r"C[ÓO]DIGO\s+DA\s+INFRA[ÇC][ÃA]O\s*[:\-]?\s*([0-9]{3,4}[\-\s]?[0-9]{1,2})",
                  txt, flags=re.IGNORECASE)
    if m:
        raw = m.group(1)
        raw = raw.replace(" ", "")
        return normalize_codigo(raw)
    # fallback: dentro do bloco de Identificação da Infração
    bloco = re.search(r"IDENTIFICA[ÇC][ÃA]O\s+DA\s+INFRA[ÇC][ÃA]O(.{0,150})",
                      txt, flags=re.IGNORECASE|re.DOTALL)
    if bloco:
        m2 = re.search(r"([0-9]{3,4}[\-\s]?[0-9]{1,2})", bloco.group(1))
        if m2:
            return normalize_codigo(m2.group(1))
    # último recurso: 'INFRAÇÃO' perto do número
    m3 = re.search(r"INFRA[ÇC][ÃA]O.{0,40}([0-9]{3,4}[\-\s]?[0-9]{1,2})",
                   txt, flags=re.IGNORECASE|re.DOTALL)
    if m3:
        return normalize_codigo(m3.group(1))
    return None

def extract_observacoes_auto(texto: str) -> str:
    """Tenta pegar OBS/OBSERVAÇÃO/OBSERVAÇÕES; se não achar, tenta frases curtas típicas como fallback."""
    upper = texto.upper()
    # 1) âncora OBS...
    m = re.search(r"\bOBS(?:ERVA[ÇC][ÃA]O(?:ES)?)?\b\s*:?", upper)
    if m:
        start = m.start()
        tail = texto[start:]
        stop_pat = r"|".join([re.escape(t) for t in NA_TITLES])
        stop = re.search(rf"(?:\n|\r|\r\n)\s*(?:{stop_pat})\b", tail, flags=re.IGNORECASE)
        bloco = tail[:stop.start()] if stop else tail
        bloco = re.sub(r"^(?is).*?OBS(?:ERVA[ÇC][ÃA]O(?:ES)?)?\s*:?", "", bloco).strip()
        bloco = norm_spaces(bloco)
        if bloco:
            # limitar em quebra de frase se capturou muito
            corte = re.search(r"(.{0,600}[\.!?])", bloco)
            return (corte.group(1).strip() if corte else bloco[:600]).strip()

    # 2) fallback: padrões curtos geralmente usados como 'observações' no NA digital
    padroes_curts = [
        r"CONDUTOR\s+N[ÃA]O\s+HABILITADO",
        r"RECUSA[OU]\s+BAF[ÔO]METRO",
        r"N[ÃA]O\s+APRESENTOU\s+DOCUMENTO",
        r"N[ÃA]O\s+DISP[ÓO]NIVEL|N[ÃA]O\s+DISPON[IÍ]VEL",
    ]
    for p in padroes_curts:
        mm = re.search(p, upper)
        if mm:
            frag = texto[mm.start():mm.end()]
            return norm_spaces(frag)

    return "(Campo de Observações não encontrado)"

def find_mbft_file(codigo: str, pasta="fichas_mbft") -> str | None:
    """Arquivos MBFT padronizados como xxx-xx*.pdf (mas aceitamos sem hífen por robustez)."""
    if not codigo or not os.path.exists(pasta):
        return None
    target = codigo.lower()
    nohy = codigo.replace("-", "").lower()
    for f in os.listdir(pasta):
        if not f.lower().endswith(".pdf"): 
            continue
        name = f.lower()
        if target in name or nohy in name:
            return os.path.join(pasta, f)
    return None

def extract_mbft_observation_context(full_text: str) -> dict:
    """Extrai seção principal 'Exemplos do Campo de Observações do AIT' + contextos 'observa' + obrigatoriedade."""
    t = full_text
    low = t.lower()
    principal = ""
    m = re.search(
        r"(exemplos\s+do\s+campo\s+de\s+observa[çc][ãa]o(?:es)?\s+do\s+ait.*?)(?:\n[A-Z][A-Z \t/()ºª0-9\.\-]{5,}\n|quando\s+autuar|defini[çc][õo]es|$)",
        low, flags=re.DOTALL|re.IGNORECASE
    )
    if m:
        principal = m.group(1)
    contextos = re.findall(r".{0,120}observa.{0,220}", low)
    obrigatorio = any(p in low for p in ["deve constar", "obrigat", "necessário", "registrar no campo"])

    def clean(s): 
        return norm_spaces(re.sub(r"\s+", " ", s or "")).strip()

    return {
        "trecho_principal": clean(principal),
        "contextos": [clean(c) for c in contextos[:12]],
        "obrigatorio": obrigatorio
    }

def similarity(a: str, b: str) -> float:
    return SequenceMatcher(None, (a or "").lower(), (b or "").lower()).ratio()

def compare_observations(obs_auto: str, mbft_ctx: dict) -> tuple[str, str, float]:
    if not obs_auto or "não encontrado" in obs_auto.lower():
        return ("❌ Observações não encontradas no Auto", "err", 0.0)
    candidates = []
    if mbft_ctx.get("trecho_principal"):
        candidates.append(mbft_ctx["trecho_principal"])
    candidates.extend(mbft_ctx.get("contextos", []))

    best = 0.0
    for c in candidates:
        best = max(best, similarity(obs_auto, c))

    # bônus por palavras-chave do principal presentes no auto
    bonus = 0.0
    principal = mbft_ctx.get("trecho_principal", "")
    if principal:
        keys = [w for w in re.findall(r"[a-zà-ú\-]{5,}", principal.lower())
                if w not in {"observações","observacao","quando","autuar","definições","procedimentos"}]
        hit = sum(1 for k in keys if k in (obs_auto or "").lower())
        if len(keys) > 0:
            bonus = min(0.15, hit / max(10, len(keys)) * 0.15)
    score = min(1.0, best + bonus)

    if score >= 0.72:  return ("✅ Condizente com a ficha MBFT", "ok", score)
    if score >= 0.45:  return ("⚠️ Parcialmente coerente (pode estar incompleto)", "warn", score)
    return ("❌ Divergente do que a ficha MBFT exige", "err", score)

# ============ QQROC BASIC ============

def qqroc_quem(texto: str) -> str:
    org = re.search(r"ÓRG[ÃA]O\s+AUTUADOR\s*[:\n]\s*(.+)", texto, re.IGNORECASE)
    return norm_spaces(org.group(1)) if org else "(Órgão/Autoridade não identificado)"

def qqroc_que(texto: str, codigo: str | None) -> str:
    desc = re.search(r"DESCRI[ÇC][ÃA]O\s+DA\s+INFRA[ÇC][ÃA]O\s*[:\n]\s*(.+)", texto, re.IGNORECASE)
    d = norm_spaces(desc.group(1)) if desc else "(Descrição não localizada)"
    return f"Código: {codigo or '—'} • {d}"

def qqroc_requisitos(texto: str) -> list[str]:
    problemas = []
    if not re.search(r"\bLOCAL DA INFRA[ÇC][ÃA]O\b", texto, re.IGNORECASE): problemas.append("Local da infração ausente")
    if not re.search(r"\bDATA\b", texto, re.IGNORECASE): problemas.append("Data ausente")
    if not re.search(r"\bHORA\b", texto, re.IGNORECASE): problemas.append("Hora ausente")
    if not re.search(r"\bPLACA\b", texto, re.IGNORECASE): problemas.append("Placa ausente")
    if re.search(r"INSTRUMENTO DE AFERI[ÇC][ÃA]O", texto, re.IGNORECASE) and re.search(r"n[aã]o dispon[íi]vel", texto, re.IGNORECASE):
        problemas.append("Instrumento de aferição 'Não disponível'")
    return problemas

def qqroc_consequencia(status_obs: str, obrigatorio: bool) -> tuple[str, str]:
    if "❌" in status_obs and obrigatorio:
        return ("Provável nulidade por descumprimento do requisito descritivo (MBFT).", "err")
    if "⚠️" in status_obs and obrigatorio:
        return ("Aparente omissão descritiva; recomendada impugnação por insuficiência de relato.", "warn")
    return ("Requisito descritivo atendido quanto ao MBFT (verificar demais vícios formais/materiais).", "ok")

# ============ DOCX (DEFESA) ============

def gerar_defesa_docx(dados: dict) -> bytes:
    """
    Gera um DOCX simples com diagnóstico e campos editáveis.
    dados: {codigo, obs_auto, status_obs, sim, obrigatorio, orgao, desc_infracao}
    """
    doc = Document()
    styles = doc.styles["Normal"].font
    styles.name = "Calibri"
    styles.size = Pt(11)

    doc.add_heading("Recurso Administrativo – Auditoria Babix (QQROC + MBFT)", level=1)
    p = doc.add_paragraph()
    p.add_run("Código da Infração: ").bold = True
    p.add_run(dados.get("codigo") or "—")

    p = doc.add_paragraph()
    p.add_run("Descrição da Infração: ").bold = True
    p.add_run(dados.get("desc_infracao") or "—")

    p = doc.add_paragraph()
    p.add_run("Órgão Autuador: ").bold = True
    p.add_run(dados.get("orgao") or "—")

    doc.add_heading("Campo de Observações do Auto", level=2)
    doc.add_paragraph(dados.get("obs_auto") or "—")

    doc.add_heading("Conclusão (MBFT)", level=2)
    doc.add_paragraph(f"Resultado: {dados.get('status_obs')} (similaridade {dados.get('sim'):.0%})")
    doc.add_paragraph("Obrigatoriedade de observações na ficha: " + ("Sim" if dados.get("obrigatorio") else "Não identificado"))

    doc.add_heading("Fundamentos (Res. 918/2022 + MBFT)", level=2)
    doc.add_paragraph("A ficha MBFT correspondente exige a descrição do comportamento observado no campo de Observações do AIT. "
                      "Constatada divergência/insuficiência, requer-se a anulação do auto por vício material/insuficiência descritiva.")

    doc.add_heading("Pedidos", level=2)
    doc.add_paragraph("a) Reconhecimento da nulidade do AIT por inobservância do padrão descritivo do MBFT;")
    doc.add_paragraph("b) Subsidiariamente, o cancelamento da penalidade por insuficiência do relato fático;")
    doc.add_paragraph("c) Notificações no endereço do requerente.")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ============ UI / PIPELINE ============

st.markdown("### 📄 Envie o Auto de Infração (PDF)")
auto_pdf = st.file_uploader("Arraste o PDF aqui (NA/SENATRAN)", type=["pdf"])

if auto_pdf:
    auto_txt = read_pdf_text(auto_pdf)
    st.markdown("<div class='card'><b>📎 Arquivo recebido</b></div>", unsafe_allow_html=True)

    codigo = extract_codigo_infracao(auto_txt)
    st.markdown(f"<div class='card'><b>🆔 CÓDIGO DA INFRAÇÃO:</b> <span class='code-pill'>{codigo or 'não localizado'}</span></div>",
                unsafe_allow_html=True)

    obs_auto = extract_observacoes_auto(auto_txt)
    st.markdown("<div class='card'><b>📝 Campo de Observações (Auto):</b>"
                f"<div class='obs' style='margin-top:.5rem;'>{(obs_auto if len(obs_auto)<1200 else obs_auto[:1200]+' …')}</div></div>",
                unsafe_allow_html=True)

    # localizar ficha MBFT
    ficha_path = find_mbft_file(codigo, "fichas_mbft")
    if not ficha_path:
        st.markdown("<div class='card'><span class='badge err'>Ficha MBFT não encontrada</span> "
                    "Coloque o PDF correspondente em <code>./fichas_mbft/</code> (ex.: <code>527-41.pdf</code>).</div>",
                    unsafe_allow_html=True)
        st.stop()

    st.markdown(f"<div class='card'>📘 Ficha MBFT: <span class='code-pill'>{os.path.basename(ficha_path)}</span></div>",
                unsafe_allow_html=True)
    with open(ficha_path, "rb") as f:
        ficha_txt = read_pdf_text(f)

    mbft_ctx = extract_mbft_observation_context(ficha_txt)
    status_obs, color_obs, score = compare_observations(obs_auto, mbft_ctx)
    st.markdown(f"<div class='card'><b>🔎 Resultado (Observações × MBFT):</b> "
                f"<span class='badge {color_obs}'>{status_obs}</span> "
                f"<span class='code-pill'>similaridade {score:.0%}</span></div>",
                unsafe_allow_html=True)

    with st.expander("👁️ Trecho principal – MBFT (observações)"):
        st.write(mbft_ctx.get("trecho_principal") or "(Nada localizado)")
    with st.expander("🔎 Outros contextos com 'observa' na ficha"):
        ctxs = mbft_ctx.get("contextos") or []
        if ctxs:
            for i, c in enumerate(ctxs, 1):
                st.markdown(f"**{i}.** {c}")
        else:
            st.write("(Nenhum contexto adicional)")

    # ===== QQROC =====
    st.markdown("## 📊 QQROC – Diagnóstico")

    quem = qqroc_quem(auto_txt)
    st.markdown(f"<div class='card'><b>Q – Quem:</b> {quem}</div>", unsafe_allow_html=True)

    que = qqroc_que(auto_txt, codigo)
    st.markdown(f"<div class='card'><b>Q – Que:</b> {que}</div>", unsafe_allow_html=True)

    req = qqroc_requisitos(auto_txt)
    if req:
        itens = "".join([f"<li>{r}</li>" for r in req])
        st.markdown(f"<div class='card'><b>R – Requisitos:</b> "
                    f"<span class='badge warn'>atenção</span><ul>{itens}</ul></div>", unsafe_allow_html=True)
    else:
        st.markdown(f"<div class='card'><b>R – Requisitos:</b> <span class='badge ok'>ok</span></div>",
                    unsafe_allow_html=True)

    st.markdown(f"<div class='card'><b>O – Observações:</b> <span class='badge {color_obs}'>{status_obs}</span></div>",
                unsafe_allow_html=True)

    cons_txt, cons_color = qqroc_consequencia(status_obs, mbft_ctx.get("obrigatorio", False))
    st.markdown(f"<div class='card'><b>C – Consequências:</b> <span class='badge {cons_color}'>{cons_txt}</span></div>",
                unsafe_allow_html=True)

    # ===== DOCX BUTTON =====
    orgao = quem
    desc_infr = re.sub(r"^Código:\s*.*?•\s*", "", que) if "•" in que else que
    doc_bytes = gerar_defesa_docx({
        "codigo": codigo,
        "obs_auto": obs_auto,
        "status_obs": status_obs,
        "sim": score,
        "obrigatorio": mbft_ctx.get("obrigatorio", False),
        "orgao": orgao,
        "desc_infracao": desc_infr
    })
    st.download_button(
        label="📝 Gerar Defesa (DOCX)",
        data=doc_bytes,
        file_name=f"Defesa_{codigo or 'AIT'}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("Babix AI © 2025 • MG Multas — Relatório auxiliar. Revise juridicamente antes de protocolar.")
